from reportlab.pdfgen import canvas
from docx import Document
from PIL import Image, ImageDraw
import os

# Create data folder if it doesn't exist
os.makedirs("data", exist_ok=True)

# ---------- Create benefits.pdf ----------
pdf_path = "data/benefits.pdf"
c = canvas.Canvas(pdf_path)

c.setFont("Helvetica", 14)
c.drawString(50, 800, "Summary of Benefits and Coverage")

c.setFont("Helvetica", 11)

lines = [
    "Primary care visits are covered.",
    "Specialist visits require referral.",
    "Emergency services are covered.",
    "Prescription drugs are included.",
    "Preventive care is covered at no cost.",
    "Mental health services are covered.",
]

y = 770

for line in lines:
    c.drawString(50, y, line)
    y -= 20

c.save()

# ---------- Create claims_process.docx ----------
doc = Document()

doc.add_heading("Claims Process", level=1)

doc.add_paragraph("Step 1: Receive treatment.")
doc.add_paragraph("Step 2: Collect medical documents.")
doc.add_paragraph("Step 3: Complete claim form.")
doc.add_paragraph("Step 4: Submit claim.")
doc.add_paragraph("Step 5: Insurance company reviews.")
doc.add_paragraph("Step 6: Payment is processed.")

doc.save("data/claims_process.docx")

# ---------- Create enrollment.png ----------
img = Image.new("RGB", (700, 350), color="white")

draw = ImageDraw.Draw(img)

draw.text((20,20),"Enrollment Form", fill="black")
draw.text((20,70),"Name: John Smith", fill="black")
draw.text((20,110),"DOB: 10 Jan 1995", fill="black")
draw.text((20,150),"Plan: Standard Health", fill="black")
draw.text((20,190),"Phone: 9876543210", fill="black")

img.save("data/enrollment.png")

print("All sample files created successfully!")